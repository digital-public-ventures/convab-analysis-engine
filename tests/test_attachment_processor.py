"""Tests for the attachment processor module."""

from unittest.mock import patch

import pytest

from src.regs_dot_gov_exploration.attachment_processor import (
    AttachmentProcessor,
    DOCXExtractor,
    PDFExtractor,
    combine_narratives,
)


class TestPDFExtractor:
    """Tests for PDF text extraction."""

    def test_extract_from_text_pdf(self, tmp_path):
        """Test extracting text from a PDF with embedded text."""
        import fitz

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Hello, World!")
        doc.save(str(pdf_path))
        doc.close()

        extractor = PDFExtractor()
        with open(pdf_path, "rb") as f:
            content = f.read()

        text = extractor.extract(content)
        assert "Hello" in text
        assert "World" in text

    def test_extract_from_empty_pdf(self, tmp_path):
        """Test extracting text from an empty PDF."""
        import fitz

        pdf_path = tmp_path / "empty.pdf"
        doc = fitz.open()
        doc.new_page()  # Empty page
        doc.save(str(pdf_path))
        doc.close()

        extractor = PDFExtractor()
        with open(pdf_path, "rb") as f:
            content = f.read()

        text = extractor.extract(content)
        assert text.strip() == ""


class TestDOCXExtractor:
    """Tests for DOCX text extraction."""

    def test_extract_from_docx(self, tmp_path):
        """Test extracting text from a DOCX file."""
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello from DOCX!")
        doc.add_paragraph("This is another paragraph.")
        doc.save(str(docx_path))

        extractor = DOCXExtractor()
        with open(docx_path, "rb") as f:
            content = f.read()

        text = extractor.extract(content)
        assert "Hello from DOCX!" in text
        assert "another paragraph" in text


class TestAttachmentProcessor:
    """Tests for the AttachmentProcessor class."""

    def test_initialization(self):
        """Test processor initialization with custom timeout."""
        processor = AttachmentProcessor(timeout=60.0)
        assert processor.timeout == 60.0

    def test_is_url_detection(self):
        """Test URL vs file path detection."""
        processor = AttachmentProcessor()

        # URLs
        assert processor._is_url("https://example.com/doc.pdf") is True
        assert processor._is_url("http://example.com/doc.pdf") is True

        # File paths
        assert processor._is_url("/path/to/doc.pdf") is False
        assert processor._is_url("relative/path/doc.pdf") is False
        assert processor._is_url("doc.pdf") is False

    def test_detect_extension_pdf(self):
        """Test PDF file extension detection."""
        processor = AttachmentProcessor()

        assert processor._detect_extension("document.pdf") == ".pdf"
        assert processor._detect_extension("document.PDF") == ".pdf"
        assert processor._detect_extension("https://example.com/doc.pdf") == ".pdf"

    def test_detect_extension_docx(self):
        """Test DOCX file extension detection."""
        processor = AttachmentProcessor()

        assert processor._detect_extension("document.docx") == ".docx"
        assert processor._detect_extension("document.DOCX") == ".docx"

    def test_get_extractor_pdf(self):
        """Test getting PDF extractor."""
        processor = AttachmentProcessor()
        extractor = processor._get_extractor(".pdf")
        assert isinstance(extractor, PDFExtractor)

    def test_get_extractor_docx(self):
        """Test getting DOCX extractor."""
        processor = AttachmentProcessor()
        extractor = processor._get_extractor(".docx")
        assert isinstance(extractor, DOCXExtractor)

    def test_get_extractor_unsupported(self):
        """Test unsupported file type raises error."""
        processor = AttachmentProcessor()

        with pytest.raises(ValueError) as exc_info:
            processor._get_extractor(".txt")
        assert "Unsupported" in str(exc_info.value)

    def test_read_local_file(self, tmp_path):
        """Test reading local file content."""
        test_file = tmp_path / "test.txt"
        test_file.write_bytes(b"test content")

        processor = AttachmentProcessor()
        content = processor._read_file(str(test_file))

        assert content == b"test content"

    def test_read_missing_file(self):
        """Test reading non-existent file raises error."""
        processor = AttachmentProcessor()

        with pytest.raises(FileNotFoundError):
            processor._read_file("/nonexistent/path/file.pdf")

    def test_extract_text_safe_returns_error_for_unsupported_type(self):
        """Test that unsupported file types return error message."""
        processor = AttachmentProcessor()

        text, error = processor.extract_text_safe("document.txt")
        assert text is None
        assert error is not None
        assert "Unsupported" in error

    @patch.object(AttachmentProcessor, "_fetch_url")
    def test_extract_text_from_url(self, mock_fetch, tmp_path):
        """Test extracting text from a URL."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "URL content test")
        pdf_bytes = doc.tobytes()
        doc.close()

        mock_fetch.return_value = pdf_bytes

        processor = AttachmentProcessor()
        text, error = processor.extract_text_safe("https://example.com/doc.pdf")

        assert error is None
        assert "URL content test" in text
        mock_fetch.assert_called_once_with("https://example.com/doc.pdf")

    def test_extract_text_from_local_file(self, tmp_path):
        """Test extracting text from a local PDF file."""
        import fitz

        pdf_path = tmp_path / "local.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Local file content")
        doc.save(str(pdf_path))
        doc.close()

        processor = AttachmentProcessor()
        text, error = processor.extract_text_safe(str(pdf_path))

        assert error is None
        assert "Local file content" in text

    def test_process_attachments(self, tmp_path):
        """Test processing multiple attachments."""
        import fitz

        # Create two test PDFs
        pdf1 = tmp_path / "doc1.pdf"
        pdf2 = tmp_path / "doc2.pdf"

        for i, path in enumerate([pdf1, pdf2], 1):
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 50), f"Document {i} content")
            doc.save(str(path))
            doc.close()

        processor = AttachmentProcessor()
        results = processor.process_attachments([str(pdf1), str(pdf2)])

        assert len(results) == 2
        assert "Document 1 content" in results[str(pdf1)]
        assert "Document 2 content" in results[str(pdf2)]


class TestCombineNarratives:
    """Tests for the combine_narratives utility function."""

    def test_inline_only(self):
        """Test combining when there's only inline comment."""
        result = combine_narratives("This is the inline comment.", {})

        assert "[Inline Comment]" in result
        assert "This is the inline comment." in result

    def test_with_attachments(self):
        """Test combining inline comment with attachment text."""
        attachment_texts = {"https://example.com/doc.pdf": "Attachment content here"}

        result = combine_narratives("See attached.", attachment_texts)

        assert "[Inline Comment]" in result
        assert "See attached." in result
        assert "[Attachment: doc.pdf]" in result
        assert "Attachment content here" in result

    def test_with_failed_attachments(self):
        """Test combining when attachment extraction failed (None text)."""
        attachment_texts: dict[str, str | None] = {"https://example.com/doc.pdf": None}

        result = combine_narratives("See attached.", attachment_texts)

        # With None text, the attachment section is not included
        assert "[Inline Comment]" in result
        assert "See attached." in result
        # The attachment with None text should be omitted
        assert "doc.pdf" not in result

    def test_multiple_attachments(self):
        """Test combining with multiple attachments."""
        attachment_texts = {
            "https://example.com/doc1.pdf": "First document",
            "https://example.com/doc2.pdf": "Second document",
        }

        result = combine_narratives("See attachments.", attachment_texts)

        assert "doc1.pdf" in result
        assert "doc2.pdf" in result
        assert "First document" in result
        assert "Second document" in result

    def test_empty_inline_comment(self):
        """Test with empty inline comment."""
        attachment_texts = {"https://example.com/doc.pdf": "All content in attachment"}

        result = combine_narratives("", attachment_texts)

        # Empty inline comment is not included
        assert "[Inline Comment]" not in result
        assert "All content in attachment" in result

    def test_empty_result(self):
        """Test with no content (empty comment and no attachments)."""
        result = combine_narratives("", {})
        assert result == ""

    def test_whitespace_only_comment(self):
        """Test with whitespace-only comment."""
        result = combine_narratives("   ", {})
        assert result == ""
